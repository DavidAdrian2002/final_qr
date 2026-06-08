import uuid
import os

from utils.qr import generar_qr
from flask import send_file
from datetime import datetime
from docx import Document
from docx.shared import Inches
from flask import *
from utils.db import *

from werkzeug.security import (
    generate_password_hash,
    check_password_hash
)
from reportlab.platypus import (
    SimpleDocTemplate,
    Table,
    TableStyle,
    Paragraph,
    Spacer
)

from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Image
from flask import (
    Flask,
    render_template,
    request,
    redirect,
    session,
    url_for,
    flash,
    send_file
)

from werkzeug.security import (
    generate_password_hash,
    check_password_hash
)

from utils.db import init_db, get_connection

app = Flask(__name__)

init_db()

app.secret_key = "super_secret_key"

UPLOAD_FOLDER = "static/img/alumnos"

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)


asistencias_temp = {}

# =========================
# INICIO
# =========================
@app.route("/")
def home():

    if "docente_id" in session:
        return redirect("/dashboard")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT COUNT(*) as total FROM docentes")
    total = cursor.fetchone()["total"]

    conn.close()

    if total == 0:
        return redirect("/register")

    return redirect("/login")


# =========================
# REGISTRO
# =========================
@app.route("/register", methods=["GET", "POST"])
def register():

    conn = get_connection()
    cursor = conn.cursor()

    if request.method == "POST":

        nombre = request.form["nombre"]
        apellido = request.form["apellido"]
        usuario = request.form["usuario"]
        password = request.form["password"]

        # Verificar si usuario existe
        cursor.execute("""
        SELECT *
        FROM docentes
        WHERE usuario = %s
        """, (usuario,))

        existe = cursor.fetchone()

        if existe:

            flash("El usuario ya existe")

            conn.close()

            return redirect("/register")

        # Hash seguro
        password_hash = generate_password_hash(password)

        # Crear docente
        cursor.execute("""
        INSERT INTO docentes (
            nombre,
            apellido,
            usuario,
            password
        )
        VALUES (%s, %s, %s, %s)
        """, (
            nombre,
            apellido,
            usuario,
            password_hash
        ))

        conn.commit()

        conn.close()

        flash("Cuenta creada correctamente")

        return redirect("/login")

    conn.close()

    return render_template("register.html")


# =========================
# LOGIN
# =========================
@app.route("/login", methods=["GET", "POST"])
def login():

    error = None

    if request.method == "POST":

        usuario = request.form["usuario"]
        password = request.form["password"]

        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
        SELECT * FROM docentes
        WHERE usuario = %s
        """, (usuario,))

        docente = cursor.fetchone()

        conn.close()

        if docente:

            if check_password_hash(
                docente["password"],
                password
            ):

                session["docente_id"] = docente["id"]
                session["docente_nombre"] = docente["nombre"]

                return redirect("/dashboard")

        error = "Usuario o contraseña incorrectos"

    return render_template(
        "login.html",
        error=error
    )



# =========================
# ESCUELAS
# =========================
@app.route("/escuelas")
def escuelas():

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT escuelas.*
    FROM escuelas

    INNER JOIN docente_escuelas
    ON escuelas.id = docente_escuelas.escuela_id

    WHERE docente_escuelas.docente_id = %s
    """, (session["docente_id"],))

    escuelas = cursor.fetchall()

    conn.close()

    return render_template(
        "escuelas.html",
        escuelas=escuelas
    )


# =========================
# CREAR ESCUELA
# =========================
@app.route("/escuelas/nueva", methods=["GET", "POST"])
def nueva_escuela():

    if "docente_id" not in session:
        return redirect("/login")

    if request.method == "POST":

        nombre = request.form["nombre"]
        numero = request.form["numero"]
        localidad = request.form["localidad"]

        conn = get_connection()
        cursor = conn.cursor()

        # Crear escuela
        cursor.execute("""
        INSERT INTO escuelas (
            nombre,
            numero,
            localidad
        )
        VALUES (%s, %s, %s)
        RETURNING id
        """, (
            nombre,
            numero,
            localidad
        ))

        escuela_id = cursor.fetchone()["id"]

        # Asociar docente
        cursor.execute("""
        INSERT INTO docente_escuelas (
            docente_id,
            escuela_id
        )
        VALUES (%s, %s)
        """, (
            session["docente_id"],
            escuela_id
        ))

        conn.commit()
        conn.close()

        return redirect("/escuelas")

    return render_template("nueva_escuela.html")


# =========================
# ELIMINAR ESCUELA
# =========================
@app.route("/escuelas/eliminar/<int:id>")
def eliminar_escuela(id):

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    # Eliminar relación
    cursor.execute("""
    DELETE FROM docente_escuelas
    WHERE docente_id = %s
    AND escuela_id = %s
    """, (
        session["docente_id"],
        id
    ))

    conn.commit()
    conn.close()

    return redirect("/escuelas")

    # =========================
    # RELACION DOCENTE-ESCUELA
    # =========================
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS docente_escuelas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        docente_id INTEGER NOT NULL,
        escuela_id INTEGER NOT NULL,

        FOREIGN KEY(docente_id)
        REFERENCES docentes(id),

        FOREIGN KEY(escuela_id)
        REFERENCES escuelas(id)
    )
    """)

    conn.commit()
    conn.close()

# =========================
# LISTAR GRADOS
# =========================
@app.route("/grados")
def grados():

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT
        grados.id,
        grados.nombre,
        escuelas.nombre as escuela_nombre

    FROM grados

    INNER JOIN escuelas
    ON grados.escuela_id = escuelas.id

    INNER JOIN docente_escuelas
    ON escuelas.id = docente_escuelas.escuela_id

    WHERE docente_escuelas.docente_id = %s
    """, (session["docente_id"],))

    grados = cursor.fetchall()

    conn.close()

    return render_template(
        "grados.html",
        grados=grados
    )


# =========================
# NUEVO GRADO
# =========================
@app.route("/grados/nuevo", methods=["GET", "POST"])
def nuevo_grado():

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    # Obtener escuelas del docente
    cursor.execute("""
    SELECT escuelas.*
    FROM escuelas

    INNER JOIN docente_escuelas
    ON escuelas.id = docente_escuelas.escuela_id

    WHERE docente_escuelas.docente_id = %s
    """, (session["docente_id"],))

    escuelas = cursor.fetchall()

    if request.method == "POST":

        nombre = request.form["nombre"]
        escuela_id = request.form["escuela_id"]

        cursor.execute("""
        INSERT INTO grados (
            nombre,
            escuela_id
        )
        VALUES (%s, %s)
        """, (
            nombre,
            escuela_id
        ))

        conn.commit()
        conn.close()

        return redirect("/grados")

    conn.close()

    return render_template(
        "nuevo_grado.html",
        escuelas=escuelas
    )


# =========================
# ELIMINAR GRADO
# =========================
@app.route("/grados/eliminar/<int:id>")
def eliminar_grado(id):

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    DELETE FROM grados
    WHERE id = %s
    """, (id,))

    conn.commit()
    conn.close()

    return redirect("/grados")

# =========================
# ALUMNOS
# =========================
@app.route("/alumnos")
def alumnos():

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT
        alumnos.*,
        grados.nombre as grado_nombre,
        escuelas.nombre as escuela_nombre

    FROM alumnos

    INNER JOIN grados
    ON alumnos.grado_id = grados.id

    INNER JOIN escuelas
    ON grados.escuela_id = escuelas.id

    INNER JOIN docente_escuelas
    ON escuelas.id = docente_escuelas.escuela_id

    WHERE docente_escuelas.docente_id = %s
    AND alumnos.activo = 1

    ORDER BY apellido ASC
    """, (session["docente_id"],))

    alumnos = cursor.fetchall()

    conn.close()

    return render_template(
        "alumnos.html",
        alumnos=alumnos
    )


# =========================
# NUEVO ALUMNO
# =========================
@app.route("/alumnos/nuevo", methods=["GET", "POST"])
def nuevo_alumno():

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT
        grados.id,
        grados.nombre,
        escuelas.nombre as escuela_nombre

    FROM grados

    INNER JOIN escuelas
    ON grados.escuela_id = escuelas.id

    INNER JOIN docente_escuelas
    ON escuelas.id = docente_escuelas.escuela_id

    WHERE docente_escuelas.docente_id = %s
    """, (session["docente_id"],))

    grados = cursor.fetchall()

    if request.method == "POST":

        nombre = request.form["nombre"]
        apellido = request.form["apellido"]
        dni = request.form["dni"]
        sexo = request.form["sexo"]
        grado_id = request.form["grado_id"]

        qr_token = str(uuid.uuid4())

        foto = None

        # =========================
        # FOTO
        # =========================
        archivo = request.files.get("foto")

        if archivo and archivo.filename != "":

            extension = archivo.filename.split(".")[-1]

            nombre_foto = f"{uuid.uuid4()}.{extension}"

            ruta = os.path.join(
                app.config["UPLOAD_FOLDER"],
                nombre_foto
            )

            archivo.save(ruta)

            foto = nombre_foto

        cursor.execute("""
        INSERT INTO alumnos (
            nombre,
            apellido,
            dni,
            sexo,
            foto,
            qr_token,
            grado_id
        )
        VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (
            nombre,
            apellido,
            dni,
            sexo,
            foto,
            qr_token,
            grado_id
        ))

        conn.commit()
        conn.close()

        flash("Alumno creado correctamente")

        return redirect("/alumnos")

    conn.close()

    return render_template(
        "nuevo_alumno.html",
        grados=grados
    )


# =========================
# DESACTIVAR ALUMNO
# =========================
@app.route("/alumnos/desactivar/<int:id>")
def desactivar_alumno(id):

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    UPDATE alumnos
    SET activo = 0
    WHERE id = %s
    """, (id,))

    conn.commit()
    conn.close()

    flash("Alumno desactivado")

    return redirect("/alumnos")

# =========================
# VER QR
# =========================
@app.route("/alumnos/qr/<int:id>")
def ver_qr(id):

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT *
    FROM alumnos
    WHERE id = %s
    """, (id,))

    alumno = cursor.fetchone()

    conn.close()

    if not alumno:
        return "Alumno no encontrado"

    generar_qr(alumno["qr_token"])

    return render_template(
        "ver_qr.html",
        alumno=alumno
    )


# =========================
# DESCARGAR QR
# =========================
@app.route("/alumnos/qr/descargar/<int:id>")
def descargar_qr(id):

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT *
    FROM alumnos
    WHERE id = %s
    """, (id,))

    alumno = cursor.fetchone()

    conn.close()

    if not alumno:
        return "Alumno no encontrado"

    ruta = generar_qr(alumno["qr_token"])

    return send_file(
        ruta,
        as_attachment=True
    )

# =========================
# TOMAR ASISTENCIA
# =========================
@app.route("/asistencia")
def asistencia():

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT
        grados.id,
        grados.nombre,
        escuelas.nombre as escuela_nombre

    FROM grados

    INNER JOIN escuelas
    ON grados.escuela_id = escuelas.id

    INNER JOIN docente_escuelas
    ON escuelas.id = docente_escuelas.escuela_id

    WHERE docente_escuelas.docente_id = %s
    """, (session["docente_id"],))

    grados = cursor.fetchall()

    conn.close()

    return render_template(
        "asistencia.html",
        grados=grados
    )


# =========================
# INICIAR ASISTENCIA
# =========================
@app.route("/asistencia/iniciar", methods=["POST"])
def iniciar_asistencia():

    if "docente_id" not in session:
        return redirect("/login")

    grado_id = str(request.form["grado_id"])

    asistencias_temp[grado_id] = []

    return redirect(f"/asistencia/escanear/{grado_id}")


# =========================
# PANTALLA ESCANEO
# =========================
@app.route("/asistencia/escanear/<int:grado_id>")
def escanear_asistencia(grado_id):

    if "docente_id" not in session:
        return redirect("/login")

    return render_template(
        "escanear.html",
        grado_id=grado_id
    )


# =========================
# REGISTRAR ESCANEO
# =========================
@app.route("/asistencia/registrar", methods=["POST"])
def registrar_asistencia():

    if "docente_id" not in session:
        return {"error": "No autorizado"}, 401

    token = request.form["token"]
    grado_id = str(request.form["grado_id"])
    grado_id_int = int(grado_id)
    estado = request.form["estado"]

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT *
    FROM alumnos
    WHERE qr_token = %s
    AND grado_id = %s
    """, (
        token,
        grado_id_int
    ))

    alumno = cursor.fetchone()

    if not alumno:
        conn.close()
        return {"error": "Alumno no encontrado"}

    # Evitar duplicados
    for item in asistencias_temp.get(grado_id, []):

        if item["alumno_id"] == alumno["id"]:
            conn.close()
            return {"error": "Alumno ya registrado"}

    hora = datetime.now().strftime("%H:%M:%S")

    data = {
        "alumno_id": alumno["id"],
        "nombre": alumno["nombre"],
        "apellido": alumno["apellido"],
        "sexo": alumno["sexo"],
        "estado": estado,
        "hora": hora
    }

    asistencias_temp[grado_id].append(data)

    conn.close()

    return data


# =========================
# CERRAR ASISTENCIA
# =========================
@app.route("/asistencia/cerrar/<int:grado_id>")
def cerrar_asistencia(grado_id):

    if "docente_id" not in session:
        return redirect("/login")

    # Normalizar tipos
    grado_id_int = int(grado_id)
    grado_id_str = str(grado_id)

    conn = get_connection()
    cursor = conn.cursor()

    fecha = datetime.now().strftime("%Y-%m-%d")

    # Obtener registros temporales
    registrados = asistencias_temp.get(grado_id_str, [])

    registrados_ids = []

    # =========================
    # GUARDAR PRESENTES/DEMORAS
    # =========================
    for item in registrados:

        registrados_ids.append(item["alumno_id"])

        cursor.execute("""
        INSERT INTO asistencias (
            alumno_id,
            fecha,
            hora,
            estado,
            grado_id
        )
        VALUES (%s, %s, %s, %s, %s)
        """, (
            item["alumno_id"],
            fecha,
            item["hora"],
            item["estado"],
            grado_id_int
        ))

    # =========================
    # OBTENER TODOS LOS ALUMNOS
    # =========================
    cursor.execute("""
    SELECT *
    FROM alumnos
    WHERE grado_id = %s
    AND activo = 1
    """, (grado_id_int,))

    alumnos = cursor.fetchall()

    # =========================
    # REGISTRAR AUSENTES
    # =========================
    for alumno in alumnos:

        if alumno["id"] not in registrados_ids:

            cursor.execute("""
            INSERT INTO asistencias (
                alumno_id,
                fecha,
                hora,
                estado,
                grado_id
            )
            VALUES (%s, %s, %s, %s, %s)
            """, (
                alumno["id"],
                fecha,
                "--:--",
                "Ausente",
                grado_id_int
            ))

    conn.commit()
    conn.close()

    # Limpiar memoria temporal
    asistencias_temp.pop(grado_id_str, None)

    flash("Asistencia guardada correctamente")

    return redirect("/dashboard")

# =========================
# HISTORIAL
# =========================
@app.route("/historial", methods=["GET"])
def historial():

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    # Obtener grados
    cursor.execute("""
    SELECT
        grados.id,
        grados.nombre,
        escuelas.nombre as escuela_nombre

    FROM grados

    INNER JOIN escuelas
    ON grados.escuela_id = escuelas.id

    INNER JOIN docente_escuelas
    ON escuelas.id = docente_escuelas.escuela_id

    WHERE docente_escuelas.docente_id = %s
    """, (session["docente_id"],))

    grados = cursor.fetchall()

    grado_id = request.args.get("grado_id")
    fecha = request.args.get("fecha")

    asistencias = []

    if grado_id and fecha:

        cursor.execute("""
        SELECT
            asistencias.*,

            alumnos.nombre,
            alumnos.apellido,
            alumnos.sexo,

            grados.nombre as grado_nombre

        FROM asistencias

        INNER JOIN alumnos
        ON asistencias.alumno_id = alumnos.id

        INNER JOIN grados
        ON asistencias.grado_id = grados.id

        WHERE asistencias.grado_id = %s
        AND asistencias.fecha = %s

        ORDER BY alumnos.apellido ASC
        """, (
            grado_id,
            fecha
        ))

        asistencias = cursor.fetchall()

    conn.close()

    return render_template(
        "historial.html",
        grados=grados,
        asistencias=asistencias
    )
# =========================
# GENERAR QR MASIVO
# =========================
@app.route("/qr")
def qr_grados():

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT
        grados.id,
        grados.nombre,
        escuelas.nombre as escuela_nombre

    FROM grados

    INNER JOIN escuelas
    ON grados.escuela_id = escuelas.id

    INNER JOIN docente_escuelas
    ON escuelas.id = docente_escuelas.escuela_id

    WHERE docente_escuelas.docente_id = %s
    """, (session["docente_id"],))

    grados = cursor.fetchall()

    conn.close()

    return render_template(
        "qr_grados.html",
        grados=grados
    )
# =========================
# VER QR POR GRADO
# =========================
@app.route("/qr/grado/<int:grado_id>")
def qr_por_grado(grado_id):

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT
        alumnos.*,
        grados.nombre as grado_nombre

    FROM alumnos

    INNER JOIN grados
    ON alumnos.grado_id = grados.id

    WHERE alumnos.grado_id = %s
    AND alumnos.activo = 1

    ORDER BY alumnos.apellido ASC
    """, (grado_id,))

    alumnos = cursor.fetchall()

    conn.close()

    # Generar QR si no existen
    for alumno in alumnos:
        generar_qr(alumno["qr_token"])

    return render_template(
        "qr_grado.html",
        alumnos=alumnos
    )
# =========================
# PDF QR MASIVO
# =========================
@app.route("/qr/pdf/<int:grado_id>")
def qr_pdf(grado_id):

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT *
    FROM alumnos
    WHERE grado_id = %s
    AND activo = 1

    ORDER BY apellido ASC
    """, (grado_id,))

    alumnos = cursor.fetchall()

    conn.close()

    pdf_path = f"qr_grado_{grado_id}.pdf"

    doc = SimpleDocTemplate(pdf_path)

    elementos = []

    styles = getSampleStyleSheet()

    titulo = Paragraph(
        "QR de Alumnos",
        styles["Title"]
    )

    elementos.append(titulo)

    elementos.append(Spacer(1, 20))

    datos = []

    fila = []

    contador = 0

    for alumno in alumnos:

        ruta_qr = generar_qr(alumno["qr_token"])

        qr_img = Image(
            ruta_qr,
            width=120,
            height=120
        )

        texto = Paragraph(
            f"""
            {alumno['apellido']},
            {alumno['nombre']}
            """,
            styles["BodyText"]
        )

        contenido = [qr_img, texto]

        fila.append(contenido)

        contador += 1

        if contador % 3 == 0:

            datos.append(fila)

            fila = []

    if fila:
        datos.append(fila)

    tabla = Table(datos)

    tabla.setStyle(TableStyle([
        ("ALIGN", (0,0), (-1,-1), "CENTER"),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("BOTTOMPADDING", (0,0), (-1,-1), 20)
    ]))

    elementos.append(tabla)

    doc.build(elementos)

    return send_file(
        pdf_path,
        as_attachment=True
    )
# =========================
# EXPORTAR PDF
# =========================
@app.route("/historial/pdf")
def historial_pdf():

    if "docente_id" not in session:
        return redirect("/login")

    grado_id = request.args.get("grado_id")
    fecha = request.args.get("fecha")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT
        asistencias.estado,
        asistencias.hora,

        alumnos.nombre,
        alumnos.apellido,

        grados.nombre as grado_nombre,
        escuelas.nombre as escuela_nombre

    FROM asistencias

    INNER JOIN alumnos
    ON asistencias.alumno_id = alumnos.id

    INNER JOIN grados
    ON asistencias.grado_id = grados.id

    INNER JOIN escuelas
    ON grados.escuela_id = escuelas.id

    WHERE asistencias.grado_id = %s
    AND asistencias.fecha = %s

    ORDER BY alumnos.apellido ASC
    """, (
        grado_id,
        fecha
    ))

    registros = cursor.fetchall()

    conn.close()

    pdf_path = "reporte_asistencia.pdf"

    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter
    )

    styles = getSampleStyleSheet()

    elementos = []

    titulo = Paragraph(
        f"Reporte de Asistencia - {fecha}",
        styles["Title"]
    )

    elementos.append(titulo)

    elementos.append(Spacer(1, 20))

    datos = [[
        "Alumno",
        "Estado",
        "Hora"
    ]]

    for item in registros:

        datos.append([
            f"{item['apellido']}, {item['nombre']}",
            item["estado"],
            item["hora"]
        ])

    tabla = Table(datos)

    tabla.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.grey),
        ("TEXTCOLOR", (0,0), (-1,0), colors.whitesmoke),

        ("GRID", (0,0), (-1,-1), 1, colors.black),

        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),

        ("BOTTOMPADDING", (0,0), (-1,0), 12)
    ]))

    elementos.append(tabla)

    doc.build(elementos)

    return send_file(
        pdf_path,
        as_attachment=True
    )
# =========================
# REPORTES
# =========================
@app.route("/reportes", methods=["GET", "POST"])
def reportes():

    if "docente_id" not in session:
        return redirect("/login")

    conn = get_connection()
    cursor = conn.cursor()

    # Obtener grados del docente
    cursor.execute("""
    SELECT
        grados.id,
        grados.nombre,
        escuelas.nombre as escuela_nombre

    FROM grados

    INNER JOIN escuelas
    ON grados.escuela_id = escuelas.id

    INNER JOIN docente_escuelas
    ON escuelas.id = docente_escuelas.escuela_id

    WHERE docente_escuelas.docente_id = %s
    """, (session["docente_id"],))

    grados = cursor.fetchall()

    conn.close()

    return render_template(
        "reportes.html",
        grados=grados
    )
# =========================
# GENERAR REPORTE
# =========================
@app.route("/reportes/generar", methods=["POST"])
def generar_reporte():

    if "docente_id" not in session:
        return redirect("/login")

    grado_id = request.form["grado_id"]
    mes = request.form["mes"]
    dias_habiles = int(
        request.form["dias_habiles"]
    )

    conn = get_connection()
    cursor = conn.cursor()

    # Obtener alumnos
    cursor.execute("""
    SELECT *
    FROM alumnos
    WHERE grado_id = %s
    AND activo = 1
    ORDER BY apellido ASC
    """, (grado_id,))

    alumnos = cursor.fetchall()

    reporte = []

    for alumno in alumnos:

        # Presentes
        cursor.execute("""
        SELECT COUNT(*) as total
        FROM asistencias
        WHERE alumno_id = %s
        AND estado = 'Presente'
        AND fecha LIKE %s
        """, (
            alumno["id"],
            f"{mes}%"
        ))

        presentes = cursor.fetchone()["total"]

        # Demoras
        cursor.execute("""
        SELECT COUNT(*) as total
        FROM asistencias
        WHERE alumno_id = %s
        AND estado = 'Demora'
        AND fecha LIKE %s
        """, (
            alumno["id"],
            f"{mes}%"
        ))

        demoras = cursor.fetchone()["total"]

        # Ausentes
        cursor.execute("""
        SELECT COUNT(*) as total
        FROM asistencias
        WHERE alumno_id = %s
        AND estado = 'Ausente'
        AND fecha LIKE %s
        """, (
            alumno["id"],
            f"{mes}%"
        ))

        ausentes = cursor.fetchone()["total"]

        porcentaje = round(
            (
                (presentes + demoras)
                / dias_habiles
            ) * 100,
            1
        )

        reporte.append({

            "apellido": alumno["apellido"],
            "nombre": alumno["nombre"],
            "presentes": presentes,
            "demoras": demoras,
            "ausentes": ausentes,
            "porcentaje": porcentaje

        })

    conn.close()

    return render_template(
        "reporte_resultado.html",
        reporte=reporte,
        mes=mes,
        dias_habiles=dias_habiles
    )
# =========================
# EXPORTAR REPORTE WORD
# =========================
@app.route("/reportes/word", methods=["POST"])
def reporte_word():

    if "docente_id" not in session:
        return redirect("/login")

    grado_nombre = request.form["grado_nombre"]
    escuela_nombre = request.form["escuela_nombre"]
    mes = request.form["mes"]
    dias_habiles = int(
        request.form["dias_habiles"]
    )

    conn = get_connection()
    cursor = conn.cursor()

    grado_id = request.form["grado_id"]

    
    cursor.execute("""
    SELECT *
    FROM alumnos
    WHERE grado_id = %s
    AND activo = 1
    ORDER BY apellido ASC
    """, (grado_id,))

    alumnos = cursor.fetchall()

    # =========================
    # CREAR WORD
    # =========================
    doc = Document()

    # TITULO
    titulo = doc.add_heading(
        "REPORTE MENSUAL DE ASISTENCIA",
        level=1
    )

    doc.add_paragraph(
        f"{escuela_nombre}"
    )

    doc.add_paragraph(
        f"{grado_nombre}"
    )

    doc.add_paragraph(
        f"Mes: {mes}"
    )

    # INTRO
    doc.add_paragraph(
        "Según los datos registrados "
        "en el sistema de asistencia, "
        "se obtuvieron los siguientes "
        "porcentajes correspondientes "
        "al período seleccionado."
    )

    # TABLA
    table = doc.add_table(
        rows=1,
        cols=5
    )

    table.style = "Table Grid"

    hdr = table.rows[0].cells

    hdr[0].text = "Alumno"
    hdr[1].text = "Presentes"
    hdr[2].text = "Ausentes"
    hdr[3].text = "Demoras"
    hdr[4].text = "%"

    for alumno in alumnos:

        # Presentes
        cursor.execute("""
        SELECT COUNT(*) as total
        FROM asistencias
        WHERE alumno_id = %s
        AND estado = 'Presente'
        AND fecha LIKE %s
        """, (
            alumno["id"],
            f"{mes}%"
        ))

        presentes = cursor.fetchone()["total"]

        # Ausentes
        cursor.execute("""
        SELECT COUNT(*) as total
        FROM asistencias
        WHERE alumno_id = %s
        AND estado = 'Ausente'
        AND fecha LIKE %s
        """, (
            alumno["id"],
            f"{mes}%"
        ))

        ausentes = cursor.fetchone()["total"]

        # Demoras
        cursor.execute("""
        SELECT COUNT(*) as total
        FROM asistencias
        WHERE alumno_id = %s
        AND estado = 'Demora'
        AND fecha LIKE %s
        """, (
            alumno["id"],
            f"{mes}%"
        ))

        demoras = cursor.fetchone()["total"]

        porcentaje = round(
            (
                (presentes + demoras)
                / dias_habiles
            ) * 100,
            1
        )

        row = table.add_row().cells

        row[0].text = (
            f"{alumno['apellido']}, "
            f"{alumno['nombre']}"
        )

        row[1].text = str(presentes)
        row[2].text = str(ausentes)
        row[3].text = str(demoras)
        row[4].text = f"{porcentaje}%"

    # CIERRE
    doc.add_paragraph("")
    doc.add_paragraph(
        "Los datos reflejan el seguimiento "
        "mensual de asistencia de los alumnos."
    )

    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph(
        "__________________________"
    )

    doc.add_paragraph(
        "Firma del docente"
    )

    # GUARDAR
    filename = (
        f"reporte_{grado_id}_{mes}.docx"
    )

    filepath = os.path.join(
        "static/reportes",
        filename
    )

    doc.save(filepath)

    conn.close()

    return redirect(
        f"/static/reportes/{filename}"
    )
# =========================
# DASHBOARD
# =========================
@app.route("/dashboard")
def dashboard():

    if "docente_id" not in session:
        return redirect("/login")

    return render_template(
        "dashboard.html",
        nombre=session["docente_nombre"]
    )


# =========================
# LOGOUT
# =========================
@app.route("/logout")
def logout():

    session.clear()

    return redirect("/login")


if __name__ == "__main__":

    app.run(
        host="0.0.0.0",
        port=5000,
        debug=True,
        use_reloader=False
    )